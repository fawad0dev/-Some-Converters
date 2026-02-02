"""Document converter module"""
import os
from pathlib import Path

def convert(input_file, output_file, target_format):
    """
    Convert documents between formats
    Supports: PDF ↔ DOCX, PDF → Images, TXT
    """
    try:
        input_ext = Path(input_file).suffix.lower().lstrip('.')
        target_format = target_format.lower()
        
        # PDF to DOCX
        if input_ext == 'pdf' and target_format == 'docx':
            return pdf_to_docx(input_file, output_file)
        
        # DOCX to PDF
        elif input_ext == 'docx' and target_format == 'pdf':
            return docx_to_pdf(input_file, output_file)
        
        # PDF to images
        elif input_ext == 'pdf' and target_format in ['png', 'jpg', 'jpeg']:
            return pdf_to_image(input_file, output_file, target_format)
        
        # Text conversions
        elif input_ext == 'txt' or target_format == 'txt':
            return text_convert(input_file, output_file)
        
        else:
            print(f"Unsupported document conversion: {input_ext} to {target_format}")
            return False
    
    except Exception as e:
        print(f"Document conversion error: {e}")
        return False

def pdf_to_docx(input_file, output_file):
    """Convert PDF to DOCX using pypdf and python-docx"""
    try:
        from pypdf import PdfReader
        from docx import Document
        
        # Read PDF
        reader = PdfReader(input_file)
        doc = Document()
        
        # Extract text from each page
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                doc.add_paragraph(text)
        
        # Save DOCX
        doc.save(output_file)
        return True
    except Exception as e:
        print(f"PDF to DOCX error: {e}")
        return False

def docx_to_pdf(input_file, output_file):
    """Convert DOCX to PDF"""
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        
        # Read DOCX
        doc = Document(input_file)
        
        # Create PDF
        c = canvas.Canvas(output_file, pagesize=letter)
        width, height = letter
        
        y_position = height - inch
        
        for para in doc.paragraphs:
            text = para.text
            if text.strip():
                # Simple text wrapping
                if y_position < inch:
                    c.showPage()
                    y_position = height - inch
                
                c.drawString(inch, y_position, text[:100])  # Basic implementation
                y_position -= 0.3 * inch
        
        c.save()
        return True
    except Exception as e:
        print(f"DOCX to PDF error: {e}")
        return False

def pdf_to_image(input_file, output_file, target_format):
    """Convert PDF to images (all pages or first page only) - always returns list of files"""
    try:
        from pdf2image import convert_from_path
        
        # Convert all pages to images
        images = convert_from_path(input_file)
        
        if not images:
            return False
        
        output_dir = os.path.dirname(output_file)
        base_name = os.path.splitext(os.path.basename(output_file))[0]
        
        output_files = []
        
        # If only one page, save with original filename
        if len(images) == 1:
            images[0].save(output_file, target_format.upper())
            output_files.append(output_file)
        else:
            # Multiple pages - save each page as separate file
            for i, image in enumerate(images, 1):
                page_output = os.path.join(output_dir, f"{base_name}_page_{i}.{target_format}")
                image.save(page_output, target_format.upper())
                output_files.append(page_output)
        
        # Return list of output files for batch handling
        return output_files
        
    except ImportError:
        # Fallback without pdf2image
        print("pdf2image not available, using basic conversion")
        return False
    except Exception as e:
        print(f"PDF to image error: {e}")
        return False

def text_convert(input_file, output_file):
    """Simple text file conversion"""
    try:
        with open(input_file, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return True
    except Exception as e:
        print(f"Text conversion error: {e}")
        return False
